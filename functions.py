import os
import json
from markdown2 import markdown
from docx import Document
from bs4 import BeautifulSoup

JSON_EXTENSION = '.json'
HEADINGS = ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']
LIST_TAG = 'ul'
LIST_ITEM_TAG = 'li'

def load_json_files(directory):
    json_files = []
    for filename in os.listdir(directory):
        if filename.endswith(JSON_EXTENSION):
            try:
                with open(os.path.join(directory, filename), 'r') as file:
                    json_files.append(json.load(file))
            except (IOError, json.JSONDecodeError) as e:
                print(f"Error loading {filename}: {e}")
    return json_files

def extract_keywords(json_data):
    return [keyword for control in json_data.get('controls', []) for keyword in control.get('keywords', [])]

def search_keywords(json_files, search_terms):
    matches = []
    for json_data in json_files:
        prefix = json_data.get('prefix', '')
        name = json_data.get('name', '')
        for control in json_data.get('controls', []):
            control_keywords = control.get('keywords', [])
            if any(term.lower() in keyword.lower() for term in search_terms for keyword in control_keywords):
                control_with_prefix = control.copy()
                control_with_prefix['id'] = f"{prefix} {control['id']}"
                matches.append((name, control_with_prefix))
                break
    return matches

def keyword_search(directory, search_terms):
    json_files = load_json_files(directory)
    return search_keywords(json_files, search_terms)

def convertMDtoWord(markdown_text):
    html_text = markdown(markdown_text)
    doc = Document()
    soup = BeautifulSoup(html_text, 'html.parser')

    for element in soup.find_all(HEADINGS + [LIST_TAG]):
        if element.name.startswith('h'):
            level = int(element.name[1]) - 1
            doc.add_heading(element.text, level)
        elif element.name == LIST_TAG:
            for li in element.find_all(LIST_ITEM_TAG):
                doc.add_paragraph(li.text, style='List Bullet')

    return doc
